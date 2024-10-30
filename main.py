from random import randint
from docx import Document
from docx.shared import Pt
from docx.shared import Mm
from datetime import datetime
import PySimpleGUI as psg
from docx import WD_ALIGN_PARAGRAPH

var = 1


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

key = Document()
style = key.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

def add_in_docx(a):
    paragraph = doc.add_paragraph()
    paragraph.add_run(a)
    fmt = paragraph.paragraph_format
    fmt.space_before = Mm(0)
    fmt.space_after = Mm(0)


def check_k(a):
    if a < 0:
        if a == -1:
            a = "- "
        else:
            a = "- " + str(abs(a))
    else:
        if a == 1:
            a = "+ "
        else:
            a = "+ " + str(a)
    return a

def add_float_num(a, b):
    if a%b == 0:
        return str(int(a/b))
    if a == 0 and b != 0:
        return "0"
    for i in range(2, max(abs(b)+1, abs(a)+1)):
        while abs(b) % i == 0 and abs(a) % i == 0:
            b = b // i
            a = a // i



    if a < 0 and b < 0 or a > 0 and b < 0:
        a *= -1
        b *= -1
    return str(int(a))+"/"+str(int(b))

def check_num_by_zero(a, b):
    num = randint(a, b)
    while num == 0:
        num = randint(a, b)
    return num

def check_d(d):
    for i in range (0, d):
        if i*i == d:
            return True
    return False

def make_window():
    variants = [i for i in range(1, 1001)]
    s1 = psg.Spin(variants, initial_value=1, readonly=True, size=3, enable_events=True, key='-VAR-')
    layout = [
        [psg.Text('Введите количество вариантов (до 1000)'), s1],
        [psg.OK(), psg.Text("", key='-OUT-')]
    ]

    window = psg.Window('Контрольная работа. Квадратные уравнения.', layout, font='_ 18', size=(700, 100))

    while True:
        event, values = window.read()
        if event == psg.WIN_CLOSED or event == 'Exit':
            break
        if event == 'OK':
            var_ = values['-VAR-']
            var = var_
            break

    window.close()

    count = 1

    for count in range(1, var + 1):
        paragraph = doc.add_paragraph()
        paragraph.add_run('Контрольная работа\nКвадратные уравнения\nВариант ' + str(count) + "\n").bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_key = key.add_paragraph()
        p_key.add_run('Вариант ' + str(count) + "\n")
        p_key.alignment = WD_ALIGN_PARAGRAPH.CENTER

        a = check_num_by_zero(-20, 20)
        b = check_num_by_zero(-20, 20)
        b_text = check_k(b)

        add_in_docx("Решите уравнения: \n№1\na) " + str(a) + "x\u00B2 " + b_text + "x = 0")

        p_key = key.add_paragraph()
        p_key.add_run("№1\na) 0; " + add_float_num(-b, a) + "\n")

        a = randint(2, 10)
        b = randint(2, 10)

        add_in_docx("б) " + str(a * a) + "x\u00B2 - " + str(b * b) + " = 0")
        p_key.add_run("б)  " + add_float_num(-b, a) + "; " + add_float_num(b, a) + "\n")

        a = check_num_by_zero(-10, 10)
        b = check_num_by_zero(-10, 10)
        c = check_num_by_zero(-10, 10)

        d = b * b - 4 * a * c
        while d < 0 or not(check_d(d)):
            b = check_num_by_zero(-10, 10)
            c = check_num_by_zero(-10, 10)
            d = b * b - 4 * a * c

        x1_numerator = int(-b + int(d ** 0.5))
        x2_numerator = int(-b - int(d ** 0.5))

        a_text = ''
        if a == -1:
            a_text = "-"
        elif a == 1:
            a_text = ""
        else:
            a_text = str(a)

        b = check_k(b)
        c = check_k(c)

        add_in_docx("\n№2\n" + a_text + "x\u00B2 " + b + " x " + c + " = 0")
        p_key.add_run("№2  " + add_float_num(x1_numerator, 2 * a) + "; " + add_float_num(x2_numerator, 2 * a) + "\n")

        x1 = check_num_by_zero(-10, 10)
        x2 = check_num_by_zero(-10, 10)
        while -1 * (x1 + x2) == 0:
            x1 = check_num_by_zero(-10, 10)
        p = -1 * (x1 + x2)
        q = x1 * x2
        if p == 1:
            p = "+ "
        else:
            p = check_k(p)
        if q == 1:
            q = "+ "
        else:
            q = check_k(q)
        add_in_docx("\n№3 Найдите подбором корни квадратного уравнения, используя теорему Виета:")
        add_in_docx("x\u00B2 " + p + "x " + q + " = 0")

        p_key.add_run("№3  " + str(x1) + "; " + str(x2) + "\n")

        x1 = check_num_by_zero(-10, 10)
        x2 = check_num_by_zero(-10, 10)
        while -1 * (x1 + x2) == 0:
            x1 = check_num_by_zero(-10, 10)

        p = -1 * (x1 + x2)
        q = x1 * x2
        if p == 1:
            p = ""
        else:
            p = check_k(p)
        if q == 1:
            q = ""
        else:
            q = check_k(q)

        if count % 2 == 0:
            add_in_docx("\n№4 Один из корней уравнения x\u00B2 " + p + "x + a = 0 равен " + str(x1)
                        + ". Найдите другой корень и коэффициент а.")
            p_key.add_run("№4  x = " + str(x2) + "; a = " + str(x1 * x2) + "\n")
        else:
            add_in_docx("\n№4 Один из корней уравнения x\u00B2 + ax " + q + " = 0 равен " + str(x1)
                        + ". Найдите другой корень и коэффициент а.")
            p_key.add_run("№4  x = " + str(x2) + "; a = " + str(-1 * (x1 + x2)) + "\n")

        a = randint(2, 10)
        b = randint(2, 10)
        add_in_docx("\n№5 Периметр прямоугольника равен " + str(2 * (a + b)) + " см, а его площадь - "
                    + str(a * b) + " см\u00B2. Найдите длины сторон прямоугольника.\n")
        p_key.add_run("№5   " + str(a) + "; " + str(b))
        doc.add_page_break()



    current_date = datetime.now()
    current_time = str(current_date.hour) + "h" + str(current_date.minute) + "m" + str(current_date.second) + "s"
    doc.save("КР Квадратные уравнения " + str(current_date.date()) + " " + current_time + ".docx")
    key.save("Ключи КР Квадратные уравнения " + str(current_date.date()) + " " + current_time + ".docx")



















#var = int(input("Введите количество вариантов: "))


w = make_window()

